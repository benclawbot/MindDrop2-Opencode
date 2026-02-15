"""
MindDrop User Guide Generator
Creates a comprehensive DOCX user guide for MindDrop
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_element(k, v):
    """Create a XML element with specified attribute"""
    element = OxmlElement(k)
    element.set(qn('a:v'), v)
    return element

def set_cell_shading(cell, color):
    """Set cell background color"""
    cell._tc.get_or_add_tcPr().append(create_element('a:shd', color))

def add_page_break(document):
    """Add a page break to the document"""
    document.add_page_break()

def add_image_placeholder(document, width=6, caption=""):
    """Add a placeholder for images with caption"""
    # Create a styled placeholder box
    placeholder_para = document.add_paragraph()
    placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder_para.paragraph_format.space_before = Pt(12)
    placeholder_para.paragraph_format.space_after = Pt(12)
    
    # Add placeholder text with styling
    run = placeholder_para.add_run("[ Screenshot Placeholder ]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(156, 163, 175)
    run.font.italic = True
    
    if caption:
        caption_para = document.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].font.size = Pt(10)
        caption_para.runs[0].font.italic = True
        caption_para.runs[0].font.color.rgb = RGBColor(100, 116, 139)

def create_minddrop_user_guide():
    """Create the MindDrop user guide document"""
    
    document = Document()
    
    # Set up styles
    styles = document.styles
    
    # Title Style
    title_style = styles['Title']
    title_style.font.size = Pt(36)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(79, 70, 229)  # Indigo
    
    # Heading 1 Style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(24)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(28, 25, 23)  # Stone 900
    
    # Heading 2 Style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(18)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(63, 63, 70)  # Stone 700
    
    # Normal Style
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(63, 63, 70)
    
    # ============================================
    # COVER PAGE
    # ============================================
    
    # Add some empty lines for spacing
    for _ in range(6):
        document.add_paragraph()
    
    # Main title
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MindDrop")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    # Subtitle
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Smart To-Do List with AI")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    # Tagline
    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Elite Productivity Workspace")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(148, 163, 184)
    
    # Add spacing
    for _ in range(8):
        document.add_paragraph()
    
    # Version info
    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("User Guide")
    run.font.size = Pt(16)
    run.font.bold = True
    
    version_info = document.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_info.add_run("Version 1.0")
    run.font.size = Pt(12)
    
    add_page_break(document)
    
    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    
    toc_title = document.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = toc_title.add_run("Table of Contents")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(28, 25, 23)
    
    document.add_paragraph()
    
    # TOC entries
    toc_entries = [
        ("1. Introduction", "3"),
        ("2. Getting Started", "3"),
        ("   2.1 Account Creation & Authentication", "3"),
        ("   2.2 Installation Options", "4"),
        ("3. Core Features", "4"),
        ("   3.1 Task Management", "4"),
        ("   3.2 AI-Powered Task Analysis", "5"),
        ("   3.3 Kanban Board View", "6"),
        ("   3.4 Timeline/Planning View", "6"),
        ("   3.5 Meeting Studio", "7"),
        ("   3.6 Focus Mode (Pomodoro)", "7"),
        ("   3.7 Voice Assistant", "8"),
        ("   3.8 AI Image Generation", "8"),
        ("4. Advanced Features", "9"),
        ("   4.1 Tags & Organization", "9"),
        ("   4.2 Keyboard Shortcuts", "9"),
        ("   4.3 Filters & Search", "10"),
        ("   4.4 Theme Customization", "10"),
        ("5. Data & Privacy", "10"),
        ("6. Tips & Best Practices", "11"),
    ]
    
    for entry, page in toc_entries:
        toc_entry = document.add_paragraph()
        toc_entry.paragraph_format.left_indent = Pt(36) if entry.startswith("   ") else Pt(18)
        run = toc_entry.add_run(entry)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(79, 70, 229)
        
        # Add dotted leader
        tab_stops = document.styles['Normal'].paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
        run.add_tab()
        run2 = toc_entry.add_run(page)
        run2.font.size = Pt(12)
    
    add_page_break(document)
    
    # ============================================
    # SECTION 1: INTRODUCTION
    # ============================================
    
    section1 = document.add_paragraph()
    run = section1.add_run("1. Introduction")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    intro_text = document.add_paragraph()
    intro_text.paragraph_format.line_spacing = 1.5
    run = intro_text.add_run("MindDrop is an intelligent task management application powered by Google's Gemini AI. It goes beyond traditional to-do lists by automatically analyzing your tasks, breaking them into actionable subtasks, suggesting priorities, and even generating unique workspace themes.")
    run.font.size = Pt(11)
    
    # Key Features Box
    document.add_paragraph()
    features_title = document.add_paragraph()
    run = features_title.add_run("Key Features")
    run.font.size = Pt(16)
    run.font.bold = True
    
    features = [
        "🤖 AI-Powered Task Analysis - Automatically prioritizes and breaks down tasks",
        "📋 Kanban Board & Timeline Views - Organize tasks your way",
        "🎤 Voice Assistant - Hands-free task creation with Gemini Live",
        "⏱️ Focus Mode - Pomodoro timer for deep work",
        "📝 Meeting Studio - Convert meeting notes into actionable tasks",
        "🎨 AI Image Generation - Create custom workspace backgrounds",
        "🏷️ Smart Tagging - Organize with hashtags and AI-suggested tags",
        "📱 Cross-Platform - Works on web and mobile (via Capacitor)"
    ]
    
    for feature in features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(feature)
        run.font.size = Pt(11)
    
    add_page_break(document)
    
    # ============================================
    # SECTION 2: GETTING STARTED
    # ============================================
    
    section2 = document.add_paragraph()
    run = section2.add_run("2. Getting Started")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    # 2.1 Authentication
    subsection2_1 = document.add_paragraph()
    run = subsection2_1.add_run("2.1 Account Creation & Authentication")
    run.font.size = Pt(18)
    run.font.bold = True
    
    auth_text = document.add_paragraph()
    auth_text.paragraph_format.line_spacing = 1.5
    run = auth_text.add_run("When you first launch MindDrop, you'll see the authentication screen. You have several options to get started:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    auth_options = [
        ("Google Sign-In", "Sign in with your Google account for seamless access"),
        ("Email/Password", "Create an account using your email address"),
        ("Continue as Guest", "Try MindDrop without an account (data stored locally)")
    ]
    
    for option, desc in auth_options:
        option_para = document.add_paragraph()
        run = option_para.add_run(f"• {option}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = option_para.add_run(desc)
        run2.font.size = Pt(11)
    
    add_image_placeholder(document, 5.5, "Figure 1: MindDrop Authentication Screen")
    
    document.add_paragraph()
    
    note_para = document.add_paragraph()
    run = note_para.add_run("Note: Guest mode stores your data locally in the browser. For cross-device access, sign in with Google or create an account.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 116, 139)
    
    # 2.2 Installation
    subsection2_2 = document.add_paragraph()
    run = subsection2_2.add_run("2.2 Installation Options")
    run.font.size = Pt(18)
    run.font.bold = True
    
    install_text = document.add_paragraph()
    install_text.paragraph_format.line_spacing = 1.5
    run = install_text.add_run("MindDrop is a web application that can also be installed as a native mobile app:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    install_options = [
        ("Web App", "Access at minddrop.app from any browser"),
        ("PWA", "Install as a Progressive Web App for offline access"),
        ("Android APK", "Build using Capacitor for native Android experience")
    ]
    
    for option, desc in install_options:
        option_para = document.add_paragraph()
        run = option_para.add_run(f"• {option}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = option_para.add_run(desc)
        run2.font.size = Pt(11)
    
    add_page_break(document)
    
    # ============================================
    # SECTION 3: CORE FEATURES
    # ============================================
    
    section3 = document.add_paragraph()
    run = section3.add_run("3. Core Features")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    # 3.1 Task Management
    subsection3_1 = document.add_paragraph()
    run = subsection3_1.add_run("3.1 Task Management")
    run.font.size = Pt(18)
    run.font.bold = True
    
    task_text = document.add_paragraph()
    task_text.paragraph_format.line_spacing = 1.5
    run = task_text.add_run("Creating tasks in MindDrop is intuitive and powerful. The AI assistant analyzes your input to provide smart suggestions:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    task_steps = [
        "1. Click the task input field or press 'N' on your keyboard",
        "2. Type your task naturally (e.g., 'Write a report about Q4 sales by Friday')",
        "3. Optionally add details or upload an image for context",
        "4. Press Enter - MindDrop will automatically:",
        "   • Analyze and set priority level (Critical, High, Medium, Low)",
        "   • Break the task into actionable subtasks",
        "   • Suggest time estimates for each subtask",
        "   • Extract relevant tags and deadlines"
    ]
    
    for step in task_steps:
        step_para = document.add_paragraph()
        run = step_para.add_run(step)
        run.font.size = Pt(11)
        if step.startswith("   •"):
            step_para.paragraph_format.left_indent = Pt(36)
    
    add_image_placeholder(document, 5.5, "Figure 2: Creating a New Task with AI Analysis")
    
    # Task Properties Table
    document.add_paragraph()
    props_title = document.add_paragraph()
    run = props_title.add_run("Task Properties")
    run.font.size = Pt(14)
    run.font.bold = True
    
    table = document.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    properties = [
        ("Property", "Description"),
        ("Title", "The main task description"),
        ("Description", "Additional context and details"),
        ("Priority", "AI-suggested: Critical, High, Medium, or Low"),
        ("Subtasks", "AI-generated actionable steps with time estimates"),
        ("Deadline", "Extracted deadline with smart reminders"),
        ("Tags", "AI-suggested and manual hashtags for organization")
    ]
    
    for i, (prop, desc) in enumerate(properties):
        cell = table.rows[i].cells[0]
        cell.text = prop
        if i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
        cell2 = table.rows[i].cells[1]
        cell2.text = desc
        cell2.paragraphs[0].runs[0].font.size = Pt(10)
    
    document.add_paragraph()
    
    # 3.2 AI Analysis
    subsection3_2 = document.add_paragraph()
    run = subsection3_2.add_run("3.2 AI-Powered Task Analysis")
    run.font.size = Pt(18)
    run.font.bold = True
    
    ai_text = document.add_paragraph()
    ai_text.paragraph_format.line_spacing = 1.5
    run = ai_text.add_run("MindDrop uses Google's Gemini AI to analyze your tasks and provide intelligent insights:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    ai_features = [
        ("Smart Prioritization", "AI evaluates urgency and importance to suggest the right priority level"),
        ("Subtask Generation", "Automatically breaks complex tasks into manageable steps"),
        ("Time Estimation", "AI estimates how long each subtask will take"),
        ("Context Awareness", "Understands deadlines, attachments, and natural language"),
        ("AI Coach", "Chat with AI about any task for guidance and suggestions")
    ]
    
    for feature, desc in ai_features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(f"{feature}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = feature_para.add_run(desc)
        run2.font.size = Pt(11)
    
    # 3.3 Kanban Board
    subsection3_3 = document.add_paragraph()
    run = subsection3_3.add_run("3.3 Kanban Board View")
    run.font.size = Pt(18)
    run.font.bold = True
    
    kanban_text = document.add_paragraph()
    kanban_text.paragraph_format.line_spacing = 1.5
    run = kanban_text.add_run("The Kanban board provides a visual overview of all your tasks organized into three columns:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    columns = [
        ("📋 To Do", "Tasks waiting to be started"),
        ("🚀 In Progress", "Tasks you're currently working on"),
        ("✅ Done", "Completed tasks")
    ]
    
    for col, desc in columns:
        col_para = document.add_paragraph()
        run = col_para.add_run(f"{col}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = col_para.add_run(desc)
        run2.font.size = Pt(11)
    
    document.add_paragraph()
    
    tip_para = document.add_paragraph()
    run = tip_para.add_run("Tip: Click on any task to open the detailed task view where you can edit, add comments, or use AI features.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    add_image_placeholder(document, 5.5, "Figure 3: Kanban Board View with Task Cards")
    
    add_page_break(document)
    
    # 3.4 Timeline View
    subsection3_4 = document.add_paragraph()
    run = subsection3_4.add_run("3.4 Timeline/Planning View")
    run.font.size = Pt(18)
    run.font.bold = True
    
    timeline_text = document.add_paragraph()
    timeline_text.paragraph_format.line_spacing = 1.5
    run = timeline_text.add_run("Switch to Timeline view (press Alt+2) for a calendar-based perspective on your tasks:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    timeline_features = [
        "📅 Day/Week/Month views for different planning horizons",
        "🖱️ Drag and drop tasks to reschedule them",
        "📊 Color-coded by priority for quick scanning",
        "➕ Click any time slot to create a new task with that deadline"
    ]
    
    for feature in timeline_features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(f"{feature}")
        run.font.size = Pt(11)
    
    # 3.5 Meeting Studio
    subsection3_5 = document.add_paragraph()
    run = subsection3_5.add_run("3.5 Meeting Studio")
    run.font.size = Pt(18)
    run.font.bold = True
    
    meeting_text = document.add_paragraph()
    meeting_text.paragraph_format.line_spacing = 1.5
    run = meeting_text.add_run("Transform meeting notes into actionable tasks automatically:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    meeting_steps = [
        "1. Switch to Meeting Studio view (Alt+3)",
        "2. Paste your meeting notes or transcription",
        "3. Click 'Process' to let AI analyze the content",
        "4. Review the structured minutes and extracted tasks",
        "5. Sync individual tasks directly to your board"
    ]
    
    for step in meeting_steps:
        step_para = document.add_paragraph()
        run = step_para.add_run(step)
        run.font.size = Pt(11)
    
    add_image_placeholder(document, 5.5, "Figure 4: Meeting Studio - Converting Notes to Tasks")
    
    # 3.6 Focus Mode
    subsection3_6 = document.add_paragraph()
    run = subsection3_6.add_run("3.6 Focus Mode (Pomodoro)")
    run.font.size = Pt(18)
    run.font.bold = True
    
    focus_text = document.add_paragraph()
    focus_text.paragraph_format.line_spacing = 1.5
    run = focus_text.add_run("Enter Focus Mode for distraction-free work sessions. Press 'F' or click the wand icon to let AI select your most important task:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    focus_features = [
        "⏱️ 25-minute work sessions with 5-minute breaks",
        "🤖 AI automatically picks your highest-priority task",
        "📝 Tick off subtasks as you complete them",
        "🔔 Browser notifications when sessions end",
        "🌙 Ambient background that adjusts to your theme"
    ]
    
    for feature in focus_features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(f"{feature}")
        run.font.size = Pt(11)
    
    document.add_paragraph()
    
    tip_para = document.add_paragraph()
    run = tip_para.add_run("Tip: Focus Mode works best with headphones on for audio cues!")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    add_image_placeholder(document, 5.5, "Figure 5: Focus Mode with Pomodoro Timer")
    
    # 3.7 Voice Assistant
    subsection3_7 = document.add_paragraph()
    run = subsection3_7.add_run("3.7 Voice Assistant")
    run.font.size = Pt(18)
    run.font.bold = True
    
    voice_text = document.add_paragraph()
    voice_text.paragraph_format.line_spacing = 1.5
    run = voice_text.add_run("Use hands-free task entry with Gemini Live voice recognition (press 'V'):")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    voice_features = [
        "🎤 Natural voice input for creating multiple tasks",
        "📝 Real-time transcription as you speak",
        "🤖 AI processes your spoken tasks automatically",
        "✅ Review and confirm extracted tasks before adding"
    ]
    
    for feature in voice_features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(f"{feature}")
        run.font.size = Pt(11)
    
    # 3.8 AI Image Generation
    subsection3_8 = document.add_paragraph()
    run = subsection3_8.add_run("3.8 AI Image Generation")
    run.font.size = Pt(18)
    run.font.bold = True
    
    image_text = document.add_paragraph()
    image_text.paragraph_format.line_spacing = 1.5
    run = image_text.add_run("Customize your workspace with AI-generated backgrounds using Gemini:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    image_steps = [
        "1. Open the Theme customization",
        "2. Describe your desired atmosphere (e.g., 'A futuristic city with neon lights')",
        "3. Click 'Generate' to create a unique 4K background",
        "4. Apply to instantly transform your workspace"
    ]
    
    for step in image_steps:
        step_para = document.add_paragraph()
        run = step_para.add_run(step)
        run.font.size = Pt(11)
    
    add_page_break(document)
    
    # ============================================
    # SECTION 4: ADVANCED FEATURES
    # ============================================
    
    section4 = document.add_paragraph()
    run = section4.add_run("4. Advanced Features")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    # 4.1 Tags & Organization
    subsection4_1 = document.add_paragraph()
    run = subsection4_1.add_run("4.1 Tags & Organization")
    run.font.size = Pt(18)
    run.font.bold = True
    
    tag_text = document.add_paragraph()
    tag_text.paragraph_format.line_spacing = 1.5
    run = tag_text.add_run("MindDrop uses a powerful tagging system to organize your tasks:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    tag_features = [
        "🏷️ Use #hashtags directly in task titles (e.g., 'Finish report #work #q4')",
        "🎨 Tags are auto-assigned colors for visual distinction",
        "🔍 Filter tasks by single or multiple tags",
        "🤖 AI suggests relevant tags based on task content"
    ]
    
    for feature in tag_features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(f"{feature}")
        run.font.size = Pt(11)
    
    # 4.2 Keyboard Shortcuts
    subsection4_2 = document.add_paragraph()
    run = subsection4_2.add_run("4.2 Keyboard Shortcuts")
    run.font.size = Pt(18)
    run.font.bold = True
    
    shortcut_text = document.add_paragraph()
    shortcut_text.paragraph_format.line_spacing = 1.5
    run = shortcut_text.add_run("Master these shortcuts to boost your productivity:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    # Shortcuts Table
    shortcuts_table = document.add_table(rows=9, cols=3)
    shortcuts_table.style = 'Table Grid'
    
    shortcut_headers = ["Shortcut", "Action", "Description"]
    for i, header in enumerate(shortcut_headers):
        cell = shortcuts_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    shortcuts = [
        ("N", "New Task", "Open task input"),
        ("V", "Voice Assistant", "Open voice input"),
        ("F", "Focus Mode", "AI-select task & start timer"),
        ("/", "Search/Filter", "Open filter panel"),
        ("Esc", "Close Modal", "Close any open modal"),
        ("Alt+1", "Board View", "Switch to Kanban board"),
        ("Alt+2", "Timeline View", "Switch to calendar view"),
        ("Alt+3", "Meeting Studio", "Switch to meeting notes")
    ]
    
    for i, (shortcut, action, desc) in enumerate(shortcuts):
        row = shortcuts_table.rows[i+1]
        row.cells[0].text = shortcut
        row.cells[1].text = action
        row.cells[2].text = desc
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    document.add_paragraph()
    
    # 4.3 Filters & Search
    subsection4_3 = document.add_paragraph()
    run = subsection4_3.add_run("4.3 Filters & Search")
    run.font.size = Pt(18)
    run.font.bold = True
    
    filter_text = document.add_paragraph()
    filter_text.paragraph_format.line_spacing = 1.5
    run = filter_text.add_run("Quickly find what you need with powerful filtering:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    filter_options = [
        "🔤 Search by task title or description",
        "🚩 Filter by priority level (Critical, High, Medium, Low)",
        "🏷️ Filter by tags",
        "📅 Filter by urgency (Today, Overdue, This Week)",
        "🔄 Combine multiple filters for precise results"
    ]
    
    for option in filter_options:
        option_para = document.add_paragraph()
        run = option_para.add_run(f"{option}")
        run.font.size = Pt(11)
    
    # 4.4 Theme Customization
    subsection4_4 = document.add_paragraph()
    run = subsection4_4.add_run("4.4 Theme Customization")
    run.font.size = Pt(18)
    run.font.bold = True
    
    theme_text = document.add_paragraph()
    theme_text.paragraph_format.line_spacing = 1.5
    run = theme_text.add_run("Personalize your MindDrop experience:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    theme_features = [
        "🌙 Dark/Light mode toggle",
        "🎨 AI-generated backgrounds with Magic Theme",
        "🖼️ Custom workspace backgrounds",
        "📱 Responsive design adapts to any screen size"
    ]
    
    for feature in theme_features:
        feature_para = document.add_paragraph()
        run = feature_para.add_run(f"{feature}")
        run.font.size = Pt(11)
    
    add_page_break(document)
    
    # ============================================
    # SECTION 5: DATA & PRIVACY
    # ============================================
    
    section5 = document.add_paragraph()
    run = section5.add_run("5. Data & Privacy")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    data_text = document.add_paragraph()
    data_text.paragraph_format.line_spacing = 1.5
    run = data_text.add_run("Your data security and privacy are our top priorities:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    data_points = [
        ("🔐 Authentication", "Secure Firebase authentication with Google or email"),
        ("💾 Data Storage", "Firestore database for cloud sync, localStorage for guests"),
        ("🤖 AI Processing", "Tasks processed by Google's Gemini AI for analysis"),
        ("📜 Data Ownership", "You own your data - export/backup anytime"),
        ("🔒 Privacy", "AI analysis happens securely - no data shared with third parties")
    ]
    
    for point, desc in data_points:
        point_para = document.add_paragraph()
        run = point_para.add_run(f"{point}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = point_para.add_run(desc)
        run2.font.size = Pt(11)
    
    document.add_paragraph()
    
    backup_para = document.add_paragraph()
    run = backup_para.add_run("Backup & Export: Your data can be exported as JSON at any time from the settings menu.")
    run.font.size = Pt(11)
    
    # ============================================
    # SECTION 6: TIPS & BEST PRACTICES
    # ============================================
    
    section6 = document.add_paragraph()
    run = section6.add_run("6. Tips & Best Practices")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    tips_text = document.add_paragraph()
    tips_text.paragraph_format.line_spacing = 1.5
    run = tips_text.add_run("Get the most out of MindDrop with these recommendations:")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    
    tips = [
        ("💡 Start Small", "Begin with 3-5 tasks to build the habit"),
        ("🎯 Daily Review", "Spend 5 minutes each morning reviewing your board"),
        ("🤖 Trust the AI", "Let Gemini suggest priorities - it learns your patterns"),
        ("⏰ Use Deadlines", "Set realistic deadlines for accountability"),
        ("🏷️ Tag Everything", "Consistent tagging makes filtering powerful"),
        ("🎤 Speak Naturally", "Voice input works best with clear, complete sentences"),
        ("🍅 Pomodoro Power", "Use Focus Mode for deep work sessions"),
        ("📝 Meeting Prep", "Use Meeting Studio right after calls while context is fresh")
    ]
    
    for tip, desc in tips:
        tip_para = document.add_paragraph()
        run = tip_para.add_run(f"{tip}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = tip_para.add_run(desc)
        run2.font.size = Pt(11)
    
    add_page_break(document)
    
    # ============================================
    # APPENDIX: QUICK REFERENCE
    # ============================================
    
    appendix = document.add_paragraph()
    run = appendix.add_run("Appendix: Quick Reference")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)
    
    document.add_paragraph()
    
    # Priority Colors
    priority_title = document.add_paragraph()
    run = priority_title.add_run("Priority Color Coding")
    run.font.size = Pt(14)
    run.font.bold = True
    
    priority_table = document.add_table(rows=5, cols=2)
    priority_table.style = 'Table Grid'
    
    priorities = [
        ("Priority", "Color"),
        ("Critical", "🔴 Red"),
        ("High", "🟠 Orange"),
        ("Medium", "🟡 Amber/Yellow"),
        ("Low", "🟢 Green")
    ]
    
    for i, (level, color) in enumerate(priorities):
        row = priority_table.rows[i]
        row.cells[0].text = level
        row.cells[1].text = color
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    document.add_paragraph()
    
    # Troubleshooting
    troubleshoot_title = document.add_paragraph()
    run = troubleshoot_title.add_run("Troubleshooting Common Issues")
    run.font.size = Pt(14)
    run.font.bold = True
    
    troubleshoot_items = [
        ("Voice not working", "Check microphone permissions in browser settings"),
        ("AI not analyzing", "Ensure you have an active internet connection"),
        ("Tasks not saving", "Try logging out and back in, or use a registered account"),
        ("Slow performance", "Clear browser cache or try incognito mode"),
        ("Missing features", "Some features require sign-in for full access")
    ]
    
    for issue, solution in troubleshoot_items:
        item_para = document.add_paragraph()
        run = item_para.add_run(f"• {issue}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run2 = item_para.add_run(solution)
        run2.font.size = Pt(11)
    
    # ============================================
    # FOOTER / CONTACT
    # ============================================
    
    add_page_break(document)
    
    contact_title = document.add_paragraph()
    contact_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_title.add_run("Need Help?")
    run.font.size = Pt(20)
    run.font.bold = True
    
    contact_text = document.add_paragraph()
    contact_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_text.add_run("For additional support, feature requests, or feedback, please reach out through our official channels.")
    run.font.size = Pt(11)
    
    document.add_paragraph()
    document.add_paragraph()
    
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("© 2024 MindDrop. All rights reserved.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(148, 163, 184)
    
    # Save the document
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'MindDrop_User_Guide.docx')
    document.save(output_path)
    print(f"User guide created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_minddrop_user_guide()
